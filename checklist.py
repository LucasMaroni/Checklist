import streamlit as st
import os
import webbrowser
import pyperclip
from datetime import datetime, timedelta
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from io import BytesIO
import smtplib
from email.message import EmailMessage
from dotenv import load_dotenv
import zipfile
import time
import pandas as pd
import re
import requests
import json
import msal

# ============ CARREGA VARIÁVEIS DE AMBIENTE PRIMEIRO ============
load_dotenv()

# Configuração da página para celular
st.set_page_config(
    page_title="Checklist de Caminhão",
    layout="centered",
    initial_sidebar_state="collapsed"  # Esconde sidebar para celular
)

# =========================================================
# CONFIGURAÇÕES
# =========================================================
SITE_ID = os.getenv("SHAREPOINT_SITE_ID", "grupotransmaroni.sharepoint.com,05ba43d5-bdf9-4038-8e14-4d1eaf3e1a6f,df8ac049-ced8-4fd7-aacd-e5e30e40892b")
LIST_ID = os.getenv("SHAREPOINT_LIST_ID", "34c630b2-37f8-45ed-8b74-5803b619e8c6")

# Configurações Azure AD
CLIENT_ID = os.getenv("AZURE_CLIENT_ID")
TENANT_ID = os.getenv("AZURE_TENANT_ID")
SCOPES = ["https://graph.microsoft.com/.default"]

# =========================================================
# NOVO: CONFIGURAÇÃO PASTA SHAREPOINT
# =========================================================
# Nome da pasta onde serão armazenados os documentos Word
SHAREPOINT_DOCS_FOLDER = "Checklists de Manutenção"
# ID do drive (pode ser obtido via API Graph)
SHAREPOINT_DRIVE_ID = os.getenv("SHAREPOINT_DRIVE_ID", "")  # Opcional: configurar no .env

# =========================================================
# FUNÇÕES DE AUTENTICAÇÃO PERSISTENTE (SIMPLIFICADA)
# =========================================================

def iniciar_autenticacao():
    """Inicia o fluxo de autenticação com Device Code"""
    try:
        app = msal.PublicClientApplication(
            client_id=CLIENT_ID,
            authority=f"https://login.microsoftonline.com/{TENANT_ID}"
        )
        
        # Iniciar fluxo de device code
        flow = app.initiate_device_flow(scopes=SCOPES)
        
        if "user_code" not in flow:
            raise ValueError("Falha ao criar fluxo de device code")
        
        return flow
    except Exception as e:
        st.error(f"Erro na autenticação: {str(e)}")
        return None

def obter_token(flow):
    """Obtém token do fluxo de autenticação"""
    try:
        app = msal.PublicClientApplication(
            client_id=CLIENT_ID,
            authority=f"https://login.microsoftonline.com/{TENANT_ID}"
        )
        
        result = app.acquire_token_by_device_flow(flow)
        
        if "access_token" in result:
            # Calcular tempo de expiração (geralmente 1 hora)
            expires_in = result.get("expires_in", 3600)
            expiry_time = datetime.now() + timedelta(seconds=expires_in)
            
            return {
                "access_token": result["access_token"],
                "expires_at": expiry_time.timestamp()
            }
        else:
            return None
    except Exception as e:
        return None

def token_valido(token_info):
    """Verifica se o token ainda é válido"""
    if not token_info:
        return False
    
    expires_at = token_info.get("expires_at", 0)
    # Considerar token válido se expira em mais de 5 minutos
    return time.time() < (expires_at - 300)

# =========================================================
# NOVAS FUNÇÕES PARA ENVIO DE DOCUMENTOS AO SHAREPOINT
# =========================================================

def criar_pasta_sharepoint(access_token, nome_pasta=SHAREPOINT_DOCS_FOLDER):
    """
    Cria a pasta no SharePoint se não existir
    Retorna o ID da pasta
    """
    try:
        # URL para criar pasta no drive do site
        if SHAREPOINT_DRIVE_ID:
            url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drives/{SHAREPOINT_DRIVE_ID}/root/children"
        else:
            # Usar drive padrão do site
            url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drive/root/children"
        
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        
        # Verificar se pasta já existe
        response = requests.get(url, headers=headers, timeout=30)
        if response.status_code == 200:
            items = response.json().get("value", [])
            for item in items:
                if item.get("name") == nome_pasta and item.get("folder"):
                    return item.get("id")  # Retorna ID da pasta existente
        
        # Se não existe, criar a pasta
        payload = {
            "name": nome_pasta,
            "folder": {},
            "@microsoft.graph.conflictBehavior": "rename"
        }
        
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        if response.status_code in [200, 201]:
            return response.json().get("id")
        else:
            st.warning(f"Não foi possível criar pasta: {response.text}")
            return None
            
    except Exception as e:
        st.warning(f"Erro ao criar pasta no SharePoint: {e}")
        return None

def enviar_documento_sharepoint(access_token, buffer_word, placa, data_hora):
    """
    Envia o documento Word para a pasta do SharePoint
    Retorna True se enviado com sucesso
    """
    try:
        # Obter ou criar ID da pasta
        folder_id = criar_pasta_sharepoint(access_token)
        if not folder_id:
            return False
        
        # Criar nome do arquivo com placa e timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        nome_arquivo = f"{placa}_{timestamp}.docx"
        
        # Remover caracteres inválidos do nome do arquivo
        nome_arquivo = re.sub(r'[<>:"/\\|?*]', '_', nome_arquivo)
        
        # URL para upload do arquivo
        if SHAREPOINT_DRIVE_ID:
            url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drives/{SHAREPOINT_DRIVE_ID}/items/{folder_id}:/{nome_arquivo}:/content"
        else:
            url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/drive/items/{folder_id}:/{nome_arquivo}:/content"
        
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        
        # Fazer upload do arquivo
        buffer_word.seek(0)  # Garantir que estamos no início do buffer
        response = requests.put(
            url, 
            headers=headers, 
            data=buffer_word.getvalue(),
            timeout=60
        )
        
        if response.status_code in [200, 201]:
            st.success(f"✅ Documento salvo no SharePoint: {nome_arquivo}")
            return True
        else:
            st.warning(f"❌ Falha ao enviar documento: {response.status_code} - {response.text}")
            return False
            
    except Exception as e:
        st.warning(f"❌ Erro ao enviar documento para SharePoint: {e}")
        return False

# =========================================================
# INTERFACE DE LOGIN SIMPLIFICADA
# =========================================================

def mostrar_tela_login():
    """Mostra tela de login simplificada"""
    st.title("📝 CheckList Manutenção")
    st.markdown("---")
    
    if st.button("🔐 Entrar com Microsoft", use_container_width=True, type="primary"):
        with st.spinner("Preparando autenticação..."):
            flow = iniciar_autenticacao()
            
            if flow:
                st.session_state.login_flow = flow
                st.session_state.user_code = flow["user_code"]
                st.session_state.verification_uri = flow["verification_uri"]
                st.rerun()
    
    st.stop()

def mostrar_tela_codigo():
    st.title("📱 Autenticação Microsoft")
    
    user_code = st.session_state.user_code
    verification_uri = st.session_state.verification_uri
    
    # Estado para feedback de cópia
    if "codigo_copiado" not in st.session_state:
        st.session_state.codigo_copiado = False
    
    # Layout em colunas para celular
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown("### 📋 **Código:**")
        
        # Área de código com seleção facilitada
        st.markdown(f"""
        <div onclick="this.select()" style="
            font-size: 24px; 
            font-weight: bold; 
            background-color: #f8f9fa; 
            padding: 20px; 
            border-radius: 10px; 
            text-align: center;
            border: 2px solid #dee2e6;
            cursor: text;
            user-select: all;
            -webkit-user-select: all;
            -moz-user-select: all;
            -ms-user-select: all;">
            {user_code}
        </div>
        """, unsafe_allow_html=True)
        
        # Botão com feedback visual
        if st.button(
            "✅ Copiado!" if st.session_state.codigo_copiado else "📋 Copiar Código",
            use_container_width=True,
            type="secondary" if not st.session_state.codigo_copiado else "primary",
            key="copy_button"
        ):
            # Tentar copiar via JavaScript
            js_code = f"""
            <script>
            // Método moderno
            if (navigator.clipboard) {{
                navigator.clipboard.writeText("{user_code}")
                    .then(() => console.log("Código copiado"))
                    .catch(err => {{
                        // Fallback para métodos antigos
                        const textArea = document.createElement("textarea");
                        textArea.value = "{user_code}";
                        document.body.appendChild(textArea);
                        textArea.select();
                        document.execCommand("copy");
                        document.body.removeChild(textArea);
                    }});
            }} else {{
                // Fallback para navegadores antigos
                const textArea = document.createElement("textarea");
                textArea.value = "{user_code}";
                document.body.appendChild(textArea);
                        textArea.select();
                        document.execCommand("copy");
                        document.body.removeChild(textArea);
            }}
            </script>
            """
            st.components.v1.html(js_code, height=0)
            st.session_state.codigo_copiado = True
            st.rerun()
        
        # Dica para copiar manualmente
        if not st.session_state.codigo_copiado:
            st.caption("💡 **Toque longo no código acima** → Selecione 'Copiar'")
    
    with col2:
        st.markdown("### 🔗 **Link:**")
        # Link grande e fácil de tocar
        st.markdown(f"""
        <a href="{verification_uri}" target="_blank" style="
            display: block;
            padding: 20px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            text-align: center;
            border-radius: 10px;
            text-decoration: none;
            font-size: 18px;
            font-weight: bold;
            margin: 10px 0;
            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            transition: transform 0.2s;
        " onmouseover="this.style.transform='scale(1.02)'" onmouseout="this.style.transform='scale(1)'">
            🌐 ABRIR PÁGINA DE LOGIN
        </a>
        """, unsafe_allow_html=True)
        
        st.caption("🔗 Toque acima para abrir em nova aba")
    
    # Instruções passo a passo
    st.markdown("---")
    
    # Container de instruções
    with st.container():
        st.markdown("### 📝 **Passo a passo:**")
        
        col_inst1, col_inst2 = st.columns(2)
        
        with col_inst1:
            st.markdown("""
            **1.** Toque no **botão azul** acima  
            **2.** Página abrirá em nova aba  
            **3.** **Toque no botão 'Copiar Código'**  
            **4.** Volte para a página aberta
            """)
        
        with col_inst2:
            st.markdown("""
            **5.** Cole o código copiado  
            **6.** Selecione sua conta da empresa  
            **7.** Aceite as permissões  
            **8.** Volte aqui e clique abaixo ↓
            """)
    
    st.markdown("---")
    
    # Botão principal com verificação
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    
    with col_btn2:
        if st.button(
            "✅ JÁ FIZ LOGIN - CONTINUAR", 
            type="primary", 
            use_container_width=True,
            disabled=not st.session_state.codigo_copiado,  # Desabilita se não copiou
            help="Copie o código primeiro para habilitar este botão"
        ):
            with st.spinner("Validando autenticação..."):
                token_info = obter_token(st.session_state.login_flow)
                if token_info:
                    st.session_state.access_token_info = token_info
                    st.session_state.autenticado = True
                    # Limpar dados temporários
                    for key in ["login_flow", "user_code", "verification_uri", "codigo_copiado"]:
                        if key in st.session_state:
                            del st.session_state[key]
                    st.rerun()
                else:
                    st.error("❌ Falha na autenticação. Verifique se:")
                    st.error("• Inseriu o código corretamente")
                    st.error("• Selecionou a conta correta")
                    st.error("• Concedeu todas as permissões")
    
    # Reset se necessário
    if st.button("🔄 Reiniciar processo", use_container_width=True):
        for key in ["login_flow", "user_code", "verification_uri", "codigo_copiado"]:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()

# =========================================================
# VERIFICAÇÃO DE AUTENTICAÇÃO
# =========================================================

# Estados da sessão
if "autenticado" not in st.session_state:
    st.session_state.autenticado = False

if "access_token_info" not in st.session_state:
    st.session_state.access_token_info = None

# Verificar autenticação
if not st.session_state.autenticado:
    # Não autenticado - mostrar tela de login
    if "login_flow" not in st.session_state:
        mostrar_tela_login()
    else:
        mostrar_tela_codigo()
else:
    # Verificar se token ainda é válido
    token_info = st.session_state.access_token_info
    
    if not token_valido(token_info):
        st.warning("⚠️ Sessão expirada. Faça login novamente.")
        st.session_state.autenticado = False
        st.rerun()

# =========================================================
# APÓS AUTENTICAÇÃO - APLICAÇÃO PRINCIPAL (INTERFACE ORIGINAL)
# =========================================================

st.title("📝 CheckList Manutenção")

# =========================================================
# FUNÇÕES DO CHECKLIST (MANTIDAS DA VERSÃO ANTERIOR)
# =========================================================

@st.cache_resource
def carregar_placas_validas():
    try:
        df_placas = pd.read_excel("placas.xlsx")
        return set(df_placas['PLACA'].str.upper().str.strip())
    except Exception as e:
        st.warning(f"Não foi possível carregar a lista de placas: {e}")
        return set()

# Carregar lista de placas válidas
PLACAS_VALIDAS = set()
try:
    df_placas = pd.read_excel("placas.xlsx")
    PLACAS_VALIDAS = set(df_placas['PLACA'].str.upper().str.strip())
except Exception as e:
    st.warning(f"Não foi possível carregar a lista de placas: {e}")

# Estados do checklist
if "etapa" not in st.session_state:
    st.session_state.etapa = 1
if "dados" not in st.session_state:
    st.session_state.dados = {}
if "imagens" not in st.session_state:
    st.session_state.imagens = []
if "fotos_nao_ok" not in st.session_state:
    st.session_state.fotos_nao_ok = {}

# -------------------
# RESPONSÁVEIS POR ITEM (grupos) - MANTIDO
# -------------------
RESPONSAVEIS = {
    (
        "ruberval.silva@transmaroni.com.br",
        "alex.franca@transmaroni.com.br",
        "jose.oliveira@transmaroni.com.br",
        "sarah.ferreira@transmaroni.com.br",
        "enielle.argolo@transmaroni.com.br",
        "michele.silva@transmaroni.com.br",
        "manutencao.frota@transmaroni.com.br",
        "sabrina.silva@transmaroni.com.br"
    ): [
        "VAZAMENTO_OLEO_MOTOR", "VAZAMENTO_AGUA_MOTOR", "OLEO_MOTOR_OK", "ARREFECIMENTO_OK",
        "OLEO_CAMBIO_OK", "OLEO_DIFERENCIAL_OK", "DIESEL_OK", "GNV_OK", "OLEO_CUBOS_OK",
        "VAZAMENTO_AR_OK", "PNEUS_OK", "PARABRISA_OK", "ILUMINACAO_OK", "FAIXAS_REFLETIVAS_OK",
        "FALHAS_PAINEL_OK"
    ],
    ("henrique.araujo@transmaroni.com.br", "amanda.soares@transmaroni.com.br", "manutencao.frota@transmaroni.com.br",): [
        "FUNCIONAMENTO_TK_OK"
    ],
    ("manutencao.frota@transmaroni.com.br", ): [
        "TACOGRAFO_OK"
    ],
    ("wesley.assumpcao@transmaroni.com.br", "manutencao.frota@transmaroni.com.br", ): [
        "FUNILARIA_OK"
    ],
    ("mirella.trindade@transmaroni.com.br", "manutencao.frota@transmaroni.com.br", ): [
        "CÂMERA_COLUNALD", "CÂMERA_COLUNALE", "CÂMERA_DEFLETORLD", "CÂMERA_DEFLETORLE",
        "CÂMERA_PARABRISA",
        "CÂMERACOLUNA_LD", "CÂMERACOLUNA_LE", "CÂMERADEFLETOR_LD", "CÂMERADEFLETOR_LE"
    ],
    (
        "manutencao.frota@transmaroni.com.br",
        "ruberval.silva@transmaroni.com.br",
        "michele.silva@transmaroni.com.br",
        "enielle.argolo@transmaroni.com.br",
        "jose.oliveira@transmaroni.com.br",
    ): [
        "PARAFUSO_SUSPENSAO_VANDERLEIA_FACCHINI"
    ],
}

def gerar_zip_imagens(imagens):
    """Cria um ZIP com as imagens da etapa 2"""
    buffer_zip = BytesIO()
    with zipfile.ZipFile(buffer_zip, "w") as zf:
        for idx, img in enumerate(imagens, start=1):
            zf.writestr(f"foto_{idx}.jpg", img.getvalue())
    buffer_zip.seek(0)
    return buffer_zip

EMAILS_OPERACOES = {
    "MERCADO - LIVRE": ["meli.operacional@transmaroni.com.br", "programacaoecommerce@transmaroni.com.br", "lucas.alves@transmaroni.com.br"],
    "BITREM": ["bitremgrupo@transmaroni.com.br"],
    "FRIGO": ["frigogrupo@transmaroni.com.br"],
    "BIMBO": ["adm.bimbo@transmaroni.com.br"],
    "BAÚ": ["baugrupo@transmaroni.com.br"]
}

def enviar_emails_personalizados(itens_nao_ok, fotos_nao_ok, checklist_itens, buffer_word, buffer_zip):
    """Envia os e-mails para os responsáveis de cada item"""
    hora_atual = datetime.now().hour
    saudacao = "Bom dia" if hora_atual < 12 else "Boa tarde"

    operacao = st.session_state.dados.get("OPERACAO", "")
    emails_operacao = EMAILS_OPERACOES.get(operacao, [])

    for destinatarios, itens_responsaveis in RESPONSAVEIS.items():
        itens_do_grupo = [i for i in itens_nao_ok if i in itens_responsaveis]
        if not itens_do_grupo:
            continue

        todos_destinatarios = list(set(destinatarios + tuple(emails_operacao)))

        msg = EmailMessage()
        msg["Subject"] = f" CHECKLIST DE MANUTENÇÃO - {st.session_state.dados.get('PLACA_CAMINHAO','')}"
        msg["From"] = os.getenv("EMAIL_USER")
        msg["To"] = ", ".join(todos_destinatarios)

        itens_texto = "\n".join([f"- {checklist_itens[i]}" for i in itens_do_grupo])
        msg.set_content(
            f"{saudacao},\n\n"
            f"Motorista: {st.session_state.dados.get('MOTORISTA','')}\n"
            f"Vistoriador: {st.session_state.dados.get('VISTORIADOR','')}\n"
            f"Data: {st.session_state.dados.get('DATA','')} {st.session_state.dados.get('HORA','')}\n\n"
            f"O veículo {st.session_state.dados.get('PLACA_CAMINHAO','')} foi verificado em seu CHECKLIST.\n"
            f"Os seguintes itens foram vistoriados e precisam ser encaminhados para manutenção:\n\n"
            f"{itens_texto}\n\n"
            "Atenciosamente,\nSistema de Checklist"
        )

        # Anexar Ficha Técnica (Word)
        msg.add_attachment(
            buffer_word.getvalue(),
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="Ficha_Tecnica.docx"
        )

        # Anexar ZIP das fotos
        msg.add_attachment(
            buffer_zip.getvalue(),
            maintype="application",
            subtype="zip",
            filename="Fotos_Checklist.zip"
        )

        # Anexar fotos dos itens NÃO OK
        for item in itens_do_grupo:
            if item in fotos_nao_ok:
                arquivos = fotos_nao_ok[item]
                if not isinstance(arquivos, list):
                    arquivos = [arquivos]
                for idx, foto in enumerate(arquivos, start=1):
                    msg.add_attachment(
                        foto.getvalue(),
                        maintype="image",
                        subtype="jpeg",
                        filename=f"{item}_{idx}.jpg"
                    )

        try:
            with smtplib.SMTP(os.getenv("EMAIL_HOST"), int(os.getenv("EMAIL_PORT"))) as smtp:
                smtp.starttls()
                smtp.login(os.getenv("EMAIL_USER"), os.getenv("EMAIL_PASS"))
                smtp.send_message(msg)
        except Exception as e:
            st.error(f"Erro ao enviar e-mail para {todos_destinatarios}: {e}")

def enviar_para_sharepoint():
    """Envia os dados do checklist para a lista do SharePoint"""
    try:
        token_info = st.session_state.access_token_info
        if not token_info:
            return False
        
        access_token = token_info["access_token"]
        
        # Mapeamento dos campos para SharePoint
        checklist_itens_mapping = {
            "OPERACAO": "field_1",
            "DATA": "field_2",
            "HORA": "field_2",
            "PLACA_CAMINHAO": "field_3",
            "PLACA_CARRETA2": "field_4",
            "TEMPO_EXECUCAO": "field_5",
            "MOTORISTA": "field_6",
            "VISTORIADOR": "field_7",
            "KM_ATUAL": "field_8",
            "TIPO_VEICULO": "field_9",
            "TIPO_CARRETA": "field_10",
            "OBSERVACOES": "field_11",
            "VAZAMENTO_OLEO_MOTOR": "field_12",
            "VAZAMENTO_AGUA_MOTOR": "field_13",
            "FUNILARIA_OK": "field_14",
            "CÂMERA_COLUNALE": "field_15",
            "CÂMERA_DEFLETORLD": "field_16",
            "CÂMERA_COLUNALD": "field_17",
            "OLEO_MOTOR_OK": "field_18",
            "TACOGRAFO_OK": "field_19",
            "FUNCIONAMENTO_TK_OK": "field_20",
            "FALHAS_PAINEL_OK": "field_21",
            "FAIXAS_REFLETIVAS_OK": "field_22",
            "ILUMINACAO_OK": "field_23",
            "PNEUS_OK": "field_24",
            "PARABRISA_OK": "field_25",
            "VAZAMENTO_AR_OK": "field_26",
            "OLEO_DIFERENCIAL_OK": "field_27",
            "OLEO_CUBOS_OK": "field_28",
            "GNV_OK": "field_29",
            "DIESEL_OK": "field_30",
            "CÂMERA_DEFLETORLE": "field_31",
            "CÂMERADEFLETOR_LE": "field_32",
            "CÂMERADEFLETOR_LD": "field_33",
            "CÂMERACOLUNA_LE": "field_34",
            "CÂMERACOLUNA_LD": "field_35",
            "CÂMERA_PARABRISA": "field_36",
            "PARAFUSO_SUSPENSAO_VANDERLEIA_FACCHINI": "field_37",
            "ARREFECIMENTO_OK": "field_38",
            "OLEO_CAMBIO_OK": "field_39",
        }
        
        # Determinar tipo de carreta
        tipo_carreta = ""
        if st.session_state.dados.get("CARRETA_2") == "X":
            tipo_carreta = "2 EIXOS"
        elif st.session_state.dados.get("CARRETA_3") == "X":
            tipo_carreta = "3 EIXOS"
        
        # Formatar DATA/HORA no formato ISO
        data_str = f"{st.session_state.dados.get('DATA', '')} {st.session_state.dados.get('HORA', '')}"
        try:
            dt_obj = datetime.strptime(data_str, "%d/%m/%Y %H:%M")
            data_iso = dt_obj.strftime("%Y-%m-%dT%H:%M:%SZ")
        except Exception:
            data_iso = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        
        # Formatar KM ATUAL como número
        km_str = str(st.session_state.dados.get('KM_ATUAL', '0')).strip()
        try:
            km_limpo = km_str.replace('.', '').replace(',', '.')
            km_numero = float(km_limpo)
        except:
            km_numero = 0
        
        # Tempo de execução
        tempo_exec = calcular_tempo_execucao()
        
        # Construir payload
        fields_data = {
            "Title": st.session_state.dados.get('PLACA_CAMINHAO', 'N/A'),
            "field_1": st.session_state.dados.get('OPERACAO', ''),
            "field_2": data_iso,
            "field_3": st.session_state.dados.get('PLACA_CAMINHAO', ''),
            "field_4": st.session_state.dados.get('PLACA_CARRETA2', ''),
            "field_5": tempo_exec,
            "field_6": st.session_state.dados.get('MOTORISTA', ''),
            "field_7": st.session_state.dados.get('VISTORIADOR', ''),
            "field_8": km_numero,
            "field_9": st.session_state.dados.get('TIPO_VEICULO', ''),
            "field_10": tipo_carreta,
            "field_11": st.session_state.dados.get('OBSERVACOES', ''),
        }
        
        # Adicionar itens do checklist
        for checklist_key, sharepoint_field in checklist_itens_mapping.items():
            if sharepoint_field.startswith('field_'):
                try:
                    field_num = int(sharepoint_field.split('_')[1])
                    if field_num >= 12:
                        status = st.session_state.dados.get(checklist_key, "")
                        if status:
                            fields_data[sharepoint_field] = status
                except:
                    continue
        
        # Remover campos vazios
        fields_data_filtrado = {k: v for k, v in fields_data.items() if v not in ["", None]}
        
        # Enviar para SharePoint
        payload_graph = {"fields": fields_data_filtrado}
        graph_url = f"https://graph.microsoft.com/v1.0/sites/{SITE_ID}/lists/{LIST_ID}/items"
        
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json"
        }
        
        response = requests.post(
            graph_url,
            headers=headers,
            json=payload_graph,
            timeout=30
        )
        
        return response.status_code in [200, 201]
            
    except Exception as e:
        return False

def calcular_tempo_execucao():
    """Calcula o tempo de execução do checklist"""
    if "start_time" in st.session_state:
        segundos = int(time.time() - st.session_state.start_time)
        minutos = segundos // 60
        segundos_restantes = segundos % 60
        return f"{minutos:02d}:{segundos_restantes:02d}"
    return "N/A"

# =========================================================
# ETAPAS DO CHECKLIST (INTERFACE ORIGINAL)
# =========================================================

# -------------------
# ETAPA 1 - DADOS DO VEÍCULO (MANTIDA ORIGINAL)
# -------------------
if st.session_state.etapa == 1:
    st.subheader("Dados do Veículo e Condutor")
    st.session_state.dados['PLACA_CAMINHAO'] = st.text_input("Placa do Caminhão", max_chars=7)

    # Inicia o temporizador quando o usuário começa a digitar
    if (
        "start_time" not in st.session_state
        and st.session_state.dados['PLACA_CAMINHAO']
    ):
        st.session_state.start_time = time.time()

    placa_digitada = st.session_state.dados['PLACA_CAMINHAO'].upper().strip()

    # Regex padrão Mercosul: LLLNLNN
    padrao_mercosul = r"^[A-Z]{3}[0-9][A-Z][0-9]{2}$"
    placa_padrao = bool(re.match(padrao_mercosul, placa_digitada))
    placa_valida = placa_digitada in PLACAS_VALIDAS if placa_digitada else False

    if placa_digitada and not placa_padrao:
        st.warning("Placa fora do padrão Mercosul! Use o formato LLLNLNN (ex: ABC1D23).")
    elif placa_digitada and not placa_valida:
        st.warning("PLACA INVÁLIDA! Verifique com sua equipe de manutenção.")

    st.session_state.dados['KM_ATUAL'] = st.text_input("KM Atual")
    st.session_state.dados['MOTORISTA'] = st.text_input("Motorista")

    # Campo OPERAÇÃO
    operacoes = [
        "MERCADO - LIVRE",
        "BITREM",
        "BIG",
        "CARREFOUR",
        "SOTREC",
        "FRIGO",
        "BIMBO",
        "UNILEVER",
        "BAÚ",
        "OUTROS",
        "PÁTIO"
    ]
    st.session_state.dados['OPERACAO'] = st.selectbox("Operação", operacoes)

    st.session_state.dados['VISTORIADOR'] = "ANTÔNIO RINALDO RAMOS"

    tipo_veiculo = st.radio("Tipo de veículo", ["CAVALO", "RÍGIDO"])
    st.session_state.dados['TIPO_VEICULO'] = tipo_veiculo

    if tipo_veiculo == "CAVALO":
        subtipo = st.radio("Configuração do Cavalo", ["TOCO 4X2", "TRUCADO 6X2", "TRAÇADO 6X4"])
        st.session_state.dados.update({
            "CAVALO_TOCO": "X" if subtipo == "TOCO 4X2" else "",
            "CAVALO_TRUCADO": "X" if subtipo == "TRUCADO 6X2" else "",
            "CAVALO_TRACADO": "X" if subtipo == "TRAÇADO 6X4" else "",
            "RIGIDO_TOCO": "",
            "RIGIDO_TRUCADO": "",
            "RIGIDO_TRACADO": "",
        })

        bitrem = st.toggle("Veículo é BITREM?")
        st.session_state.dados['BITREM'] = "SIM" if bitrem else "NÃO"
        st.session_state.dados['PLACA_CARRETA1'] = st.text_input("Placa Carreta 1", max_chars=8)

        if bitrem:
            st.session_state.dados['PLACA_CARRETA2'] = st.text_input("Placa Carreta 2", max_chars=8)
        else:
            st.session_state.dados['PLACA_CARRETA2'] = ""

        tipo_carreta = st.radio("Tipo de Carreta", ["2 EIXOS", "3 EIXOS"])
        st.session_state.dados["CARRETA_2"] = "X" if tipo_carreta == "2 EIXOS" else ""
        st.session_state.dados["CARRETA_3"] = "X" if tipo_carreta == "3 EIXOS" else ""

    else:
        subtipo = st.radio("Configuração do Rígido", ["TOCO 4X2", "TRUCADO 6X2", "TRAÇADO 6X4"])
        st.session_state.dados.update({
            "RIGIDO_TOCO": "X" if subtipo == "TOCO 4X2" else "",
            "RIGIDO_TRUCADO": "X" if subtipo == "TRUCADO 6X2" else "",
            "RIGIDO_TRACADO": "X" if subtipo == "TRAÇADO 6X4" else "",
            "CAVALO_TOCO": "",
            "CAVALO_TRUCADO": "",
            "CAVALO_TRACADO": "",
        })  
        st.session_state.dados['PLACA_CARRETA1'] = ""
        st.session_state.dados['PLACA_CARRETA2'] = ""
        st.session_state.dados['BITREM'] = "NÃO"
        st.session_state.dados["CARRETA_2"] = ""
        st.session_state.dados["CARRETA_3"] = ""

    if st.button("Avançar ➡️", disabled=not (placa_padrao and placa_valida)):
        if all([
            st.session_state.dados['PLACA_CAMINHAO'],
            st.session_state.dados['KM_ATUAL'],
            st.session_state.dados['MOTORISTA']
        ]):
            st.session_state.dados["DATA"] = datetime.now().strftime("%d/%m/%Y")
            st.session_state.dados["HORA"] = datetime.now().strftime("%H:%M")
            st.session_state.etapa = 2
        else:
            st.warning("Preencha todos os campos obrigatórios.")

# -------------------
# ETAPA 2 - FOTOS (MANTIDA ORIGINAL)
# -------------------
elif st.session_state.etapa == 2:
    st.subheader("Inserção das Imagens")
    st.image("Checklist.png", caption="Exemplo dos ângulos corretos", use_container_width=True)

    imagens = st.file_uploader(
        "Envie ao menos 4 fotos",
        type=['jpg', 'jpeg', 'png'],
        accept_multiple_files=True
    )
    if imagens and len(imagens) >= 4:
        st.session_state.imagens = imagens

    col1, col2 = st.columns(2)
    if col1.button("⬅️ Voltar"):
        st.session_state.etapa = 1
        st.rerun()
    if col2.button("Avançar ➡️"):
        if st.session_state.imagens and len(st.session_state.imagens) >= 4:
            st.session_state.etapa = 3
        else:
            st.warning("Envie no mínimo 4 imagens.")

# -------------------
# ETAPA 3 - CHECKLIST (MANTIDA ORIGINAL COM FLUXO AUTOMÁTICO)
# -------------------
elif st.session_state.etapa == 3:
    st.subheader("Etapa 3: Checklist")
    checklist_itens = {
        "ARREFECIMENTO_OK": "Nível do líquido de arrefecimento",
        "OLEO_MOTOR_OK": "Nível de óleo de motor",
        "VAZAMENTO_OLEO_MOTOR": "Vazamento de óleo motor",
        "VAZAMENTO_AGUA_MOTOR": "Vazamento de água motor",
        "OLEO_CAMBIO_OK": "Vazamento de óleo câmbio",
        "OLEO_DIFERENCIAL_OK": "Vazamento de óleo diferencial",
        "OLEO_CUBOS_OK": "Vazamento de óleo cubos",
        "DIESEL_OK": "Vazamento de diesel",
        "GNV_OK": "Vazamento de GNV",
        "VAZAMENTO_AR_OK": "Vazamento de ar",
        "PNEUS_OK": "Pneus avariados",
        "FAIXAS_REFLETIVAS_OK": "Faixas refletivas",
        "FUNILARIA_OK": "Itens avariados para funilaria",
        "ILUMINACAO_OK": "Iluminação",
        "PARABRISA_OK": "Para-brisa",
        "FALHAS_PAINEL_OK": "Presença de falhas no painel",
        "TACOGRAFO_OK": "Funcionamento tacógrafo",
        "CÂMERA_PARABRISA": "Câmera do para-brisa",
        "CÂMERA_COLUNALD": "Câmera Coluna Lado Direito",
        "CÂMERA_COLUNALE": "Câmera Coluna Lado Esquerdo",
        "CÂMERA_DEFLETORLD": "Câmera Defletor Lado Direito",
        "CÂMERA_DEFLETORLE": "Câmera Defletor Lado Esquerdo",
        "FUNCIONAMENTO_TK_OK": "Funcionamento TK",
        "CÂMERACOLUNA_LD": "Imagem digital câmera coluna LD",
        "CÂMERACOLUNA_LE": "Imagem digital câmera coluna LE",
        "CÂMERADEFLETOR_LD": "Imagem digital câmera defletor LD",
        "CÂMERADEFLETOR_LE": "Imagem digital câmera defletor LE",
        "PARAFUSO_SUSPENSAO_VANDERLEIA_FACCHINI": "Parafuso suspensão Vanderleia Facchini",
    }

    for chave, descricao in checklist_itens.items():
        opcao = st.radio(
            descricao,
            ["OK", "NÃO OK"],
            index=0,
            key=f"radio_{chave}",
            horizontal=True
        )
        st.session_state.dados[chave] = opcao
        if opcao == "NÃO OK":
            fotos = st.file_uploader(
                f"Fotos de {descricao}",
                type=['jpg', 'jpeg', 'png'],
                key=f"foto_{chave}",
                accept_multiple_files=True
            )
            if fotos:
                st.session_state.fotos_nao_ok[chave] = fotos

    # Campo OBSERVAÇÕES
    st.session_state.dados["OBSERVACOES"] = st.text_area("Observações", placeholder="Digite informações adicionais, se necessário.")

    if "finalizando" not in st.session_state:
        st.session_state.finalizando = False

    col1, col2 = st.columns(2)
    if col1.button("⬅️ Voltar"):
        st.session_state.etapa = 2
        st.rerun()

    if col2.button("✅ Finalizar Checklist", disabled=st.session_state.finalizando):
        st.session_state.finalizando = True
        with st.spinner("Finalizando checklist..."):
            try:
                # ===== Gera Word =====
                doc = Document("Ficha Técnica.docx")
                for p in doc.paragraphs:
                    for k, v in st.session_state.dados.items():
                        token = f"{{{{{k}}}}}"
                        if token in p.text:
                            p.text = p.text.replace(token, str(v))
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                for k, v in st.session_state.dados.items():
                                    token = f"{{{{{k}}}}}"
                                    if token in p.text:
                                        p.text = p.text.replace(token, str(v))

                buffer_word = BytesIO()
                doc.save(buffer_word)
                buffer_word.seek(0)

                # ZIP das fotos
                buffer_zip = gerar_zip_imagens(st.session_state.imagens)

                # Itens NÃO OK
                itens_nao_ok = [k for k, v in st.session_state.dados.items() if v == "NÃO OK"]

                # Envia e-mails
                enviar_emails_personalizados(
                    itens_nao_ok,
                    st.session_state.fotos_nao_ok,
                    checklist_itens,
                    buffer_word,
                    buffer_zip
                )
                
                # ===== NOVO: ENVIAR DOCUMENTO WORD PARA SHAREPOINT =====
                placa = st.session_state.dados.get('PLACA_CAMINHAO', '')
                data_hora = f"{st.session_state.dados.get('DATA', '')}_{st.session_state.dados.get('HORA', '').replace(':', '')}"
                
                # Criar uma cópia do buffer para enviar ao SharePoint
                buffer_word_sharepoint = BytesIO()
                buffer_word.seek(0)
                buffer_word_sharepoint.write(buffer_word.getvalue())
                buffer_word_sharepoint.seek(0)
                
                token_info = st.session_state.access_token_info
                if token_info:
                    enviado_sharepoint_doc = enviar_documento_sharepoint(
                        token_info["access_token"],
                        buffer_word_sharepoint,
                        placa,
                        data_hora
                    )
                else:
                    enviado_sharepoint_doc = False
                    st.warning("❌ Token de acesso não disponível para enviar ao SharePoint")
                
                # ENVIO PARA LISTA DO SHAREPOINT (original)
                enviou_sharepoint_lista = enviar_para_sharepoint()
                
                # Mensagens de sucesso
                mensagens = []
                if enviou_sharepoint_lista:
                    mensagens.append("✅ Checklist Registrado com Sucesso")
                if enviado_sharepoint_doc:
                    mensagens.append("✅ Documento Word Salvo")
                
                if mensagens:
                    st.success("\n".join(mensagens))
                else:
                    st.warning("⚠️ Checklist concluído, mas houve problema ao enviar para SharePoint.")

                # ⭐⭐ NOVO: AGUARDAR E VOLTAR AUTOMATICAMENTE ⭐⭐
                time.sleep(3)
                
                # LIMPA APENAS DADOS DO CHECKLIST, MANTÉM AUTENTICAÇÃO
                for key in ["etapa", "dados", "imagens", "fotos_nao_ok", "start_time", "finalizando"]:
                    if key in st.session_state:
                        del st.session_state[key]
                
                # VOLTA PARA ETAPA 1 AUTOMATICAMENTE
                st.rerun()

            except Exception as e:
                st.session_state.finalizando = False
                st.error(f"Erro ao finalizar checklist: {e}")
